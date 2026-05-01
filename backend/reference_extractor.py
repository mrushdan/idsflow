"""
Reference extractor — pulls patent reference numbers out of uploaded files.

Supports: PDF, DOCX, XLSX, XLS, CSV, plain text.

Deliberately permissive: clients send messy files (scanned PDFs,
informal email exports, Excel sheets with references mixed into
descriptive columns). The goal is to find as many valid references
as possible, then let the user review/edit the results before running
automation against them.
"""

import io
import re
from typing import List, Tuple

import pdfplumber
from docx import Document
from openpyxl import load_workbook


# Regex matches patent reference formats commonly seen in client files.
# Order matters: longer/more-specific patterns first so they aren't
# swallowed by shorter ones.
REFERENCE_PATTERNS = [
    # KR with double hyphens: KR10-2020-0123456
    re.compile(r"\bKR\s*(10-\d{4}-\d{7})\b", re.IGNORECASE),
    # WO with slash: WO2020/123456, WO 2020/123456
    re.compile(r"\bWO\s*(\d{4}[/-]\d{6})\b", re.IGNORECASE),
    # US published application: US 2020/0123456, US2020/0123456
    re.compile(r"\bUS\s*(\d{4}[/-]\d{7})\b", re.IGNORECASE),
    # JP with hyphen: JP2009-207252
    re.compile(r"\bJP\s*(\d{4}-\d{6})\b", re.IGNORECASE),
    # Standard format: country code followed by digits, no separators
    # (US8124474, EP3999999, CN101521173, etc.)
    re.compile(r"\b(US|EP|JP|CN|WO|KR|DE|FR|GB|CA|AU|IN)\s*(\d{6,12})\b", re.IGNORECASE),
]

VALID_COUNTRIES = {"US", "EP", "JP", "CN", "WO", "KR", "DE", "FR", "GB", "CA", "AU", "IN"}


def normalize_reference(raw: str) -> str:
    """Normalize a reference string to canonical format.

    Examples:
        'jp 2009-207252'   -> 'JP2009-207252'
        'US 8,124,474'     -> 'US8124474'
        'wo2020/123456'    -> 'WO2020/123456'
        'EP 3,999,999'     -> 'EP3999999'
    """
    raw = raw.strip().upper()
    raw = raw.replace(",", "")
    raw = re.sub(r"\s+", "", raw)
    return raw


def extract_from_text(text: str) -> List[str]:
    """Find patent reference numbers in free text.

    Returns a deduplicated, order-preserved list of normalized references.
    """
    # Strip commas from numeric sequences so 8,124,474 -> 8124474.
    # Run multiple times to catch chained groupings.
    text = re.sub(r"(\d),(\d)", r"\1\2", text)
    text = re.sub(r"(\d),(\d)", r"\1\2", text)
    text = re.sub(r"(\d),(\d)", r"\1\2", text)

    found = []
    seen = set()

    for pattern in REFERENCE_PATTERNS:
        for match in pattern.finditer(text):
            full = match.group(0)
            normalized = normalize_reference(full)

            country = normalized[:2]
            if country not in VALID_COUNTRIES:
                continue

            # Reject obvious false positives (years, short numbers, etc.)
            digits = re.sub(r"\D", "", normalized[2:])
            if len(digits) < 6:
                continue

            if normalized not in seen:
                seen.add(normalized)
                found.append(normalized)

    return found


def extract_from_pdf(file_bytes: bytes) -> Tuple[List[str], str]:
    """Extract references from a PDF.

    Returns (references, extraction_note). The note describes what
    was parsed — useful for telling the user when OCR might be needed.
    """
    references = []
    pages_with_text = 0
    total_pages = 0

    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            total_pages = len(pdf.pages)
            combined_text = []

            for page in pdf.pages:
                text = page.extract_text() or ""

                # Pull tables too — client IDS lists often come as tables
                for table in page.extract_tables() or []:
                    for row in table:
                        row_text = " ".join(cell or "" for cell in row)
                        text += "\n" + row_text

                if text.strip():
                    pages_with_text += 1
                combined_text.append(text)

            full_text = "\n".join(combined_text)
            references = extract_from_text(full_text)

        if pages_with_text == 0 and total_pages > 0:
            note = (
                f"PDF has {total_pages} pages but no extractable text — "
                "this looks like a scanned document. OCR would be needed."
            )
        elif not references:
            note = (
                f"Extracted text from {pages_with_text}/{total_pages} pages "
                "but found no reference numbers. Check that the file "
                "contains patent citations, or paste them manually."
            )
        else:
            note = (
                f"Found {len(references)} reference(s) in {pages_with_text}"
                f"/{total_pages} pages."
            )

        return references, note

    except Exception as e:
        return [], f"Could not read PDF: {e}"


def extract_from_docx(file_bytes: bytes) -> Tuple[List[str], str]:
    """Extract references from a Word document (.docx)."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []

        for para in doc.paragraphs:
            text_parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)

        full_text = "\n".join(text_parts)
        references = extract_from_text(full_text)

        if not references:
            note = (
                "Word document parsed but no reference numbers found. "
                "Check the content or paste manually."
            )
        else:
            note = f"Found {len(references)} reference(s) in Word document."

        return references, note

    except Exception as e:
        return [], f"Could not read Word document: {e}"


def extract_from_xlsx(file_bytes: bytes) -> Tuple[List[str], str]:
    """Extract references from an Excel file (.xlsx)."""
    try:
        wb = load_workbook(io.BytesIO(file_bytes), data_only=True, read_only=True)
        text_parts = []
        sheets_read = 0

        for sheet in wb.worksheets:
            sheets_read += 1
            for row in sheet.iter_rows(values_only=True):
                for cell in row:
                    if cell is not None:
                        text_parts.append(str(cell))

        full_text = "\n".join(text_parts)
        references = extract_from_text(full_text)

        if not references:
            note = (
                f"Read {sheets_read} sheet(s) from Excel file but no "
                "reference numbers found. Check content or paste manually."
            )
        else:
            note = (
                f"Found {len(references)} reference(s) across "
                f"{sheets_read} sheet(s) in Excel file."
            )

        return references, note

    except Exception as e:
        return [], f"Could not read Excel file: {e}"


def extract_from_csv(file_bytes: bytes) -> Tuple[List[str], str]:
    """Extract references from a CSV file. Treats it as plain text."""
    try:
        text = file_bytes.decode("utf-8", errors="replace")
        references = extract_from_text(text)

        if not references:
            note = (
                "CSV parsed but no reference numbers found. "
                "Check content or paste manually."
            )
        else:
            note = f"Found {len(references)} reference(s) in CSV."

        return references, note

    except Exception as e:
        return [], f"Could not read CSV: {e}"


def extract_references(filename: str, file_bytes: bytes) -> Tuple[List[str], str]:
    """Route to the right extractor based on file extension.

    Returns (references, extraction_note).
    """
    name_lower = filename.lower()

    if name_lower.endswith(".pdf"):
        return extract_from_pdf(file_bytes)
    elif name_lower.endswith(".docx"):
        return extract_from_docx(file_bytes)
    elif name_lower.endswith((".xlsx", ".xlsm")):
        return extract_from_xlsx(file_bytes)
    elif name_lower.endswith((".csv", ".tsv", ".txt")):
        return extract_from_csv(file_bytes)
    else:
        return [], (
            f"Unsupported file type: {filename}. "
            "Accepted formats: PDF, Word (.docx), Excel (.xlsx), CSV, plain text."
        )
